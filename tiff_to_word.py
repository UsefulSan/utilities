"""Guide https://github.com/UB-Mannheim/tesseract/wiki"""

import tkinter as tk
from tkinter import filedialog
import os
import docx
import pytesseract
from PIL import Image
from get_new_filename import get_new_filename

tess_config = '--psm 3'
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\usefu\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'


def tiff_to_docx(tiff_file, docx_file):
    # Открываем TIFF файл с помощью Pillow
    im = Image.open(tiff_file)
    # Распознаем текст на изображении с помощью pytesseract
    text = pytesseract.image_to_string(im, lang='rus', config=tess_config)
    try:
        # Открываем существующий Word-документ
        doc = docx.Document(docx_file)
    except docx.opc.exceptions.PackageNotFoundError:
        # Создаем новый Word-документ, если файл не существует
        doc = docx.Document()
    doc.add_paragraph(text)
    # Сохраняем Word документ
    doc.save(docx_file)


#
# def processing_to_folder(input_folder):
#     for file in os.listdir(input_folder):
#         if file.lower().endswith(".tiff"):
#             input_file = os.path.join(input_folder, file)
#             # base_filename = os.path.splitext(file)[0]
#             # output_file = os.path.join(input_folder, base_filename + '.docx')
#             # output_file = input_file.rjust('\\')
#             # base_filename = os.path.splitext(input_file)[0] + '.docx'
#             # # print(base_filename)
#             output_file = 'path' + ".docx"
#             tiff_to_docx(input_file, output_file)


def processing_to_folder():
    tiff_files = tiff_entry.get().split(', ')
    docx_file = docx_entry.get()
    directory = ''.join(tiff_files[0]).rsplit('/', 1)[0].replace('/', '\\')
    docx_file = directory + '\\' + docx_file + '.docx'
    print(docx_file)
    docx_file = get_new_filename(docx_file)
    print(docx_file)
    for file in tiff_files:
        if file.lower().endswith(".tiff"):
            # input_file = os.path.join(input_folder, file)
            # base_filename = os.path.splitext(file)[0]
            # output_file = os.path.join(input_folder, base_filename + '.docx')
            # output_file = input_file.rjust('\\')
            # base_filename = os.path.splitext(input_file)[0] + '.docx'
            # # print(base_filename)
            # output_file = 'path' + ".docx"
            tiff_to_docx(file, docx_file)


def select_tiff_files():
    file_paths = filedialog.askopenfilenames(filetypes=[("TIFF files", ".tiff"), ("All files", ".*")])
    tiff_entry.delete(0, tk.END)
    tiff_entry.insert(0, ', '.join(file_paths))


# def name_file_docx():
#     # Считываем текст из виджета Entry
#     input_text = entry.get()
#     # Передаем значение на бэкенд
#     process_input(input_text)


# def select_docx_folder():
#     folder_path = filedialog.askdirectory()
#     docx_entry.delete(0, tk.END)
#     docx_entry.insert(0, folder_path)


# def convert():
#     tiff_files = tiff_entry.get().split(', ')
#     docx_folder = docx_entry.get()
#     # print(docx_folder)
#     # if len(tiff_files) == 1:
#     #     pass
#     if tiff_files and docx_folder:
#         for tiff_file in tiff_files:
#             # file_name = os.path.splitext(os.path.basename(tiff_file))[0]
#             # print(os.path.splitext(os.path.basename(tiff_file)))
#
#             # docx_file = os.path.join(docx_folder, f"{file_name}.docx").replace('/', '\\')
#             docx_file = 'path'+ ".docx"
#
#             print(docx_file)
#             tiff_to_docx(tiff_file, docx_file)

app = tk.Tk()
app.title("TIFF to DOCX Converter")

tiff_label = tk.Label(app, text="TIFF файлы:")
tiff_label.grid(row=0, column=0, sticky="e")
tiff_entry = tk.Entry(app, width=60)
tiff_entry.grid(row=0, column=1)
tiff_button = tk.Button(app, text="Обзор", command=select_tiff_files)
tiff_button.grid(row=0, column=2)

docx_label = tk.Label(app, text="Имя DOCX файла:")
docx_label.grid(row=1, column=0, sticky="e")
docx_entry = tk.Entry(app, width=60)
docx_entry.grid(row=1, column=1)
# docx_button = tk.Button(app, text="Обзор", command=select_docx_folder)
# docx_button.grid(row=1, column=2)

convert_button = tk.Button(app, text="Конвертировать", command=processing_to_folder)
convert_button.grid(row=2, column=1)

app.mainloop()

# def main():
#     input_folder = 'path'
#     processing_to_folder(input_folder)
#
#
# if __name__ == '__main__':
#     main()
