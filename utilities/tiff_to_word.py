"""Guide https://github.com/UB-Mannheim/tesseract/wiki"""

import tkinter as tk
from pathlib import Path
from tkinter import filedialog

import docx
import pytesseract
from PIL import Image

from utilities.get_new_filename import get_new_filename

tess_config = '--psm 3'
# Указывается место установки tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\usefu\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'


def tiff_to_docx(tiff_file: str, docx_file: Path) -> None:
    """
    Открывает tiff файл, распознаем текст на изображении с помощью pytesseract и сохраняет его в docx
    :param tiff_file: путь к tiff файлу
    :param docx_file: путь к docx файлу
    """
    im = Image.open(tiff_file)
    text = pytesseract.image_to_string(im, lang='rus', config=tess_config)
    try:
        doc = docx.Document(str(docx_file))
    except docx.opc.exceptions.PackageNotFoundError:
        doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(docx_file)


def processing_to_folder() -> None:
    tiff_files = tiff_entry.get().split(', ')
    docx_file = docx_entry.get()
    directory = ''.join(tiff_files[0]).rsplit('/', 1)[0]
    docx_file = directory + '/' + docx_file + '.docx'
    docx_file_path = get_new_filename(docx_file)
    for file in tiff_files:
        if file.lower().endswith(".tiff"):
            tiff_to_docx(file, docx_file_path)


def select_tiff_files() -> None:
    file_paths = filedialog.askopenfilenames(filetypes=[("TIFF files", ".tiff"), ("All files", ".*")])
    tiff_entry.delete(0, tk.END)
    tiff_entry.insert(0, ', '.join(file_paths))


app = tk.Tk()
app.title("TIFF to DOCX Converter")

tiff_label = tk.Label(app, text="TIFF файлы:")
tiff_label.grid(row=0, column=0, sticky="e")
tiff_entry = tk.Entry(app, width=40)
tiff_entry.grid(row=0, column=1)
tiff_button = tk.Button(app, text="Обзор", command=select_tiff_files)
tiff_button.grid(row=0, column=2)

docx_label = tk.Label(app, text="Имя DOCX файла:")
docx_label.grid(row=1, column=0, sticky="e")
docx_entry = tk.Entry(app, width=40)
docx_entry.grid(row=1, column=1)

convert_button = tk.Button(app, text="Конвертировать", command=processing_to_folder)
convert_button.grid(row=2, column=1)

app.mainloop()
